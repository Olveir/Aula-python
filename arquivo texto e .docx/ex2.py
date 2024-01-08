from docx import Document

doc = Document()

dis = input("Qual a disciplina? ")
p = doc.add_paragraph("Disciplina:  ")
p.add_run(dis).bold = True

carga = input("Qual a carga horária? ")
p = doc.add_paragraph("Com carga horária de ")
p.add_run(carga).bold = True
p.add_run(' horas.')

#criação de tabela
tab = doc.add_table(rows=1, cols=2)
tab.style="Colorful Grid Accent 6"
cels = tab.rows[0].cells
cels[0].text = 'Numeros'
cels[1].text = "Assunto"

for numero in range(1, 5):
   assunto = input("Qual é o assunto número-{}? ".format(numero))
   dados = tab.add_row().cells
   dados[0].text = str(numero)
   dados[0].width = 5
   dados[1].text = assunto


doc.save('c:\\aula0906\\meudoc.docx')
